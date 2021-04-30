import win32com,re
import os
import docx;
import xlwt
import xlsxwriter
import shutil
from win32com.client import Dispatch, constants
from docx import Document;
from win32com import client as wc
import zip
from xlrd import open_workbook
from xlutils.copy import copy
import xlrd
import re



#workbook = xlsxwriter.Workbook(r'E:\教务处工作\19年校级质量工程\result.xls')  #生成表格
#worksheet = workbook.add_worksheet(u'文件读取记录表')   #在文件中创建文件读取记录表,如果已经读取过就不再读取
#worksheet1 = workbook.add_worksheet(u'19年校级质量工程统计总表-学院汇总表')
#worksheet2 = workbook.add_worksheet(u'结题材料中提取信息')
#worksheet3 = workbook.add_worksheet(u'延期材料中提取信息')
#worksheet.set_column('A:A',20)  #设置第一列宽度为20像素
#bold=workbook.add_format({'bold':True}) #设置一个加粗的格式对象
#file = r'E:\教务处工作\19年校级质量工程\邮件附件\临时文件\上交\张准+基于FPGA的数字电子技术课程升级\结题验收表+张准+基于FPGA的数字电子技术课程升级.doc'
#file = r'E:\教务处工作\19年校级质量工程\邮件附件\临时文件\上交\张准+基于FPGA的数字电子技术课程升级\信息光电子科技学院+张准+基于FPGA的数字电子技术课程升级结题验收表.docx'

'''
date='Fri, 26 Apr 2019 15:08:37 +0800'
file = r'E:\新建 DOC 文档.doc'
Mpath=r'E:\教务处工作\19年校级质量工程'
'''
def doSaveAas(file):
	if not os.path.exists(os.path.splitext(file)[0] + '.docx'):
		word = wc.Dispatch('Word.Application')
		try:
			doc = word.Documents.Open(file)        # 目标路径下的文件
			doc.SaveAs(os.path.splitext(file)[0]+'.docx', 12, False, "", True, "", False, False, False, False)  # 转化后路径下的文件
			doc.Close()
			word.Quit()
		except:
			'文件不存在或者打不开'


def read_word(Mpath,file,date):
	word = win32com.client.Dispatch('word.application')
	'''
	设置Word的可见性visible，默认情况下，你看不到Word程序；然后设置Word的警告信息是否出现，默认也是不出现，这样你在使用python控制Word的时候不会弹出Word的警告信息。
	'''
	word.displayalerts = 0
	word.visible = 0
	#打开读取excel，并判断此文件是否已经读取过。
	book = xlrd.open_workbook(r'E:\教务处工作\19年校级质量工程\result.xls')

	# 通过sheet_by_index()获取的sheet没有write()方法
	sheet1 = book.sheet_by_name(u'文件读取记录表')
	nrows = sheet1.nrows  # 行数
	#判断文件是否读取过
	flag_filename2excel='';
	for rownum in range(0, nrows):
		if sheet1.cell(rownum,0).value==file and sheet1.cell(rownum,1).value==date:
			return '该文件已读取'

	if '~$' in file:
		return '该文件为临时文件'
	if os.path.splitext(file)[1]=='.doc':
		doSaveAas(file)#将doc另存为docx
		try:
			os.remove(file)#删除doc文件
		except:
			print(file+'文件不存在')
		file=os.path.splitext(file)[0]+'.docx'
	try:
		doc = Document(file)  # 读入文件
	except:
		print(file,'不存在或者打不开')

	try:
		table = doc.tables[0] #获取文件中的表格集
		keywords1 = '结题验收自查汇总表'
		flag='';
		#for i in range(0, len(table.rows)):  # 从表格第一行开始循环读取表格数据
		for i in range(0, 3):  # 从表格第一行开始循环读取表格数据
			if flag=='结题验收自查汇总表':
				break
			for j in range(0, len(table.columns)):
				result = table.cell(i, j).text;
				if keywords1 in result:
					filename1=result.replace('表1：', '')
					filename1 = filename1 + date
					filename1=filename1.replace('/', '')
					filename1=filename1.replace(',', '')
					filename1 = filename1.replace(',', '')
					filename1 = filename1.replace(':', '-')
					absolute_file_path=os.path.join(os.path.split(file)[0], filename1) + '.docx'
					flag= '结题验收自查汇总表'
					break
	except:
		return '不包含表格，不是所要材料'

	if flag == '结题验收自查汇总表':
		for i in range(2, len(table.rows)):  # 从表格第二行开始循环读取表格数据
			book = xlrd.open_workbook(r'E:\教务处工作\19年校级质量工程\result.xls')
			write_book = copy(book)
			write_sheet = write_book.get_sheet(u'19年校级质量工程结题总表-来自学院汇总表')
			sheet_sta = book.sheet_by_name(u'19年校级质量工程结题总表-来自学院汇总表')
			nrows_sta = sheet_sta.nrows  # 行数
			if not re.compile(u'[\u4e00-\u9fa5]+').search(table.cell(i, 3).text):
				break
			if '省级' in table.cell(i, 10).text:
				book = xlrd.open_workbook(r'E:\教务处工作\19年校级质量工程\result.xls')
				write_book = copy(book)
				write_sheet = write_book.get_sheet(u'19年校级质量工程入选省级汇总表-来自学院汇总表')
				sheet_sta = book.sheet_by_name(u'19年校级质量工程入选省级汇总表-来自学院汇总表')
				nrows_sta = sheet_sta.nrows  # 行数
			write_sheet.write(nrows_sta, 0, table.cell(i, 0).text)
			write_sheet.write(nrows_sta, 1, table.cell(i, 1).text)
			write_sheet.write(nrows_sta, 2, table.cell(i, 3).text)
			write_sheet.write(nrows_sta, 3, table.cell(i, 4).text)
			write_sheet.write(nrows_sta, 4, table.cell(i, 5).text)
			write_sheet.write(nrows_sta, 5, table.cell(i, 6).text)
			write_sheet.write(nrows_sta, 6, table.cell(i, 7).text)
			write_sheet.write(nrows_sta, 7, table.cell(i, 8).text)
			write_sheet.write(nrows_sta, 8, table.cell(i, 9).text)
			write_sheet.write(nrows_sta, 9, table.cell(i, 10).text)
			write_sheet.write(nrows_sta, 10, date)
			#nrows_sta = nrows_sta + 1;
			write_book.save(r'E:\教务处工作\19年校级质量工程\result.xls')

	try:
		table = doc.tables[1] #获取文件中的表格集
	except:
		if flag != '':
			# 暂时不移动和重命名，一会将下面注释掉
			os.rename(file, absolute_file_path)
			# 暂时不移动，一会将下面注释掉
			shutil.move(absolute_file_path, Mpath)
			flag_filename2excel='ok'
		return '不包含延期建设项目自查汇总表'
	keywords1 = '申请延期建设项目自查汇总表'
	flag1='';
	for i in range(0, min(len(table.rows),3)):  # 从表格第一行开始循环读取表格数据
		if flag1=='申请延期建设项目自查汇总表':
			break
		for j in range(0, len(table.columns)):
			result = table.cell(i, j).text;
			if keywords1 in result:
				flag1= '申请延期建设项目自查汇总表'
				break

	if flag1 == '申请延期建设项目自查汇总表':
		book = xlrd.open_workbook(r'E:\教务处工作\19年校级质量工程\result.xls')
		write_book = copy(book)
		write_sheet = write_book.get_sheet(u'19年校级质量工程延期汇总表-来自学院汇总表')
		sheet_sta = book.sheet_by_name(u'19年校级质量工程延期汇总表-来自学院汇总表')
		nrows_sta = sheet_sta.nrows  # 行数
		for i in range(2, len(table.rows)):  # 从表格第二行开始循环读取表格数据
			if not re.compile(u'[\u4e00-\u9fa5]+').search(table.cell(i, 3).text):
				#print(re.compile(u'[\u4e00-\u9fa5]+').search(table.cell(i, 3).text))
				break
			write_sheet.write(nrows_sta, 0, table.cell(i, 0).text)
			write_sheet.write(nrows_sta, 1, table.cell(i, 1).text)
			write_sheet.write(nrows_sta, 2, table.cell(i, 3).text)
			write_sheet.write(nrows_sta, 3, table.cell(i, 4).text)
			write_sheet.write(nrows_sta, 4, table.cell(i, 5).text)
			write_sheet.write(nrows_sta, 5, table.cell(i, 6).text)
			write_sheet.write(nrows_sta, 6, table.cell(i, 7).text)
			write_sheet.write(nrows_sta, 7, table.cell(i, 8).text)
			write_sheet.write(nrows_sta, 8, table.cell(i, 9).text)
			write_sheet.write(nrows_sta, 9, table.cell(i, 10).text)
			write_sheet.write(nrows_sta, 10, table.cell(i, 11).text)
			write_sheet.write(nrows_sta, 11, date)
			nrows_sta = nrows_sta + 1;
		write_book.save(r'E:\教务处工作\19年校级质量工程\result.xls')
		if flag != '':
			# 暂时不移动和重命名，一会将下面注释掉
			os.rename(file, absolute_file_path)
			# 暂时不移动，一会将下面注释掉
			shutil.move(absolute_file_path, Mpath)
			flag_filename2excel='ok'



	#暂时先不做操作，一会将注释去掉
	txt = doc.paragraphs;
	flag = 0;
	filename2=''
	leixing=''
	for text_flag in range(0, (len(txt) - 1)):
		if '结题验收表' in txt[text_flag].text:
			leixing='结题验收表'
		if '延期申请表' in txt[text_flag].text:
			leixing='延期申请表'
		if '单    位：' in txt[text_flag].text:
			danwei=txt[text_flag].text.replace('单    位：', '')
			danwei=danwei.replace(' ', '')
		if '项目类别：' in txt[text_flag].text:
			leibie=txt[text_flag].text.replace('项目类别：', '')
			leibie=leibie.replace(' ', '')
		if '项目名称' in txt[text_flag].text:
			mingcheng=txt[text_flag].text.replace('项目名称：', '')
			mingcheng = mingcheng.replace('/', '')
			mingcheng = mingcheng.replace('\\', '')
			mingcheng = mingcheng.replace(' ', '')
			mingcheng = mingcheng.replace("\n", '')
			mingcheng = mingcheng.replace("\t", '')
		if '项目负责人：' in txt[text_flag].text:
			fuzeren=txt[text_flag].text.replace('项目负责人：', '')
			fuzeren=fuzeren.replace(' ', '')
			fuzeren=fuzeren.replace('/', '')
			fuzeren = fuzeren.replace("\t", '')
		if '立项时间：' in txt[text_flag].text:
			time_start=txt[text_flag].text.replace('立项时间：', '')
			time_start=time_start.replace(' ', '')
		if '拟结题时间：' in txt[text_flag].text:
			time_end = txt[text_flag].text.replace('拟结题时间：', '')
			time_end = time_end.replace(' ', '')
			break
		if '原结题时间：' in txt[text_flag].text:
			oraltime_end = txt[text_flag].text.replace('原结题时间：', '')
			oraltime_end = oraltime_end.replace(' ', '')
		if '延期后结题时间：' in txt[text_flag].text:
			delaytime_end = txt[text_flag].text.replace('延期后结题时间：', '')
			delaytime_end = delaytime_end.replace(' ', '')
			break

	if leixing=='结题验收表':

		table = doc.tables[0]  # 获取文件中的表格集
		for i in range(0, len(table.rows)):  # 从表格第二行开始循环读取表格数据
			for j in range(0, len(table.columns)):
				result = table.cell(i, j).text;
				if '项目类别' in result:
					table_leibie=table.cell(i, j+1).text
				if '项目名称' in result:
					table_mingcheng=table.cell(i, j+1).text
				if '项目负责人' in result:
					table_fuzeren=table.cell(i, j+1).text
					table_fuzeren=table_fuzeren.replace('/', '')
		table = doc.tables[len(doc.tables) - 1]  # 读取结题材料中的数据并写入到Excel中
		zijin=table.cell(1, 0).text;
		expertopinion = table.cell(len(table.rows)-3, 0).text;
		unitcomments=table.cell(len(table.rows) - 2, 0).text;

		book = xlrd.open_workbook(r'E:\教务处工作\19年校级质量工程\result.xls')
		write_book = copy(book)
		write_sheet = write_book.get_sheet(u'结题材料中提取信息')
		sheet2 = book.sheet_by_name(u'结题材料中提取信息')
		nrows = sheet2.nrows  # 行数
		write_sheet.write(nrows, 0, leixing)
		write_sheet.write(nrows, 1, danwei)
		write_sheet.write(nrows, 2, leibie)
		write_sheet.write(nrows, 3, mingcheng)
		write_sheet.write(nrows, 4, time_start)
		write_sheet.write(nrows, 5, time_end)
		write_sheet.write(nrows, 6, table_leibie)
		write_sheet.write(nrows, 7, table_mingcheng)
		write_sheet.write(nrows, 8, table_fuzeren)
		write_sheet.write(nrows, 9, zijin)  # 资金使用情况
		write_sheet.write(nrows, 11, expertopinion)  # 专家意见
		write_sheet.write(nrows, 12,unitcomments)  # 单位意见
		write_book.save(r'E:\教务处工作\19年校级质量工程\result.xls')
		filename2 = danwei + '+' + fuzeren + '+' + mingcheng + leixing;
		project_name=os.path.join(os.path.split(file)[0], filename2) + '.docx'
		if not os.path.isdir(project_name):#
			os.rename(file, project_name)
		project_fold_name=os.path.join(os.path.split(os.path.split(file)[0])[0], fuzeren + '+' + mingcheng)
		os.rename(os.path.split(file)[0], project_fold_name)

		project_fold_leibie=os.path.join(r'E:\教务处工作\19年校级质量工程\项目结题验收材料',leibie);

		if not os.path.isdir(project_fold_leibie):
			os.mkdir(project_fold_leibie)  # 如果没有该文件夹就创建一个
		project_fold_move=os.path.join(project_fold_leibie,os.path.split(project_fold_name)[1])
		if os.path.isdir(project_fold_move):#判断原来项目类别文件夹中是否存在该项目文件
			shutil.rmtree(project_fold_move)
		shutil.move(project_fold_name, project_fold_leibie)
		flag_filename2excel = 'ok'


	if leixing == '延期申请表':

		table = doc.tables[0]  # 获取文件中的表格集
		for i in range(0, len(table.rows)):  # 从表格第二行开始循环读取表格数据
			for j in range(0, len(table.columns)):
				result = table.cell(i, j).text;
				if '项目类别' in result:
					table_leibie = table.cell(i, j + 1).text
				if '项目名称' in result:
					table_mingcheng = table.cell(i, j + 1).text
				if '项目负责人' in result:
					table_fuzeren = table.cell(i, j + 1).text
					table_fuzeren = table_fuzeren.replace('/', '')
				if '原定结题时间' in result:
					table_oraltime = table.cell(i, j + 1).text
				if '延期后结题时间' in result:
					table_endtime = table.cell(i, j + 1).text
		table = doc.tables[len(doc.tables) - 1]  # 读取结题材料中的数据并写入到Excel中
		unitcomments = table.cell(len(table.rows) - 2, 0).text;
		book = xlrd.open_workbook(r'E:\教务处工作\19年校级质量工程\result.xls')
		write_book = copy(book)
		write_sheet = write_book.get_sheet(u'延期材料中提取信息')
		sheet2 = book.sheet_by_name(u'延期材料中提取信息')
		nrows = sheet2.nrows  # 行数
		write_sheet.write(nrows, 0, leixing)
		write_sheet.write(nrows, 1, danwei)
		write_sheet.write(nrows, 2, leibie)
		write_sheet.write(nrows, 3, mingcheng)
		write_sheet.write(nrows, 4, time_start)
		write_sheet.write(nrows, 5, oraltime_end)
		write_sheet.write(nrows, 6, delaytime_end)
		write_sheet.write(nrows, 7, table_leibie)
		write_sheet.write(nrows, 8, table_mingcheng)
		write_sheet.write(nrows, 9, table_fuzeren)
		write_sheet.write(nrows, 10, unitcomments)  # 单位意见
		write_book.save(r'E:\教务处工作\19年校级质量工程\result.xls')
		filename2 = danwei + '+' + fuzeren + '+' + mingcheng + leixing;
		project_name = os.path.join(os.path.split(file)[0], filename2) + '.docx'
		if not os.path.isdir(project_name):  #
			os.rename(file, project_name)
		project_fold_name = os.path.join(os.path.split(os.path.split(file)[0])[0], fuzeren + '+' + mingcheng)
		os.rename(os.path.split(file)[0], project_fold_name)
		project_fold_leibie = os.path.join(r'E:\教务处工作\19年校级质量工程\延期材料', leibie);
		if not os.path.isdir(project_fold_leibie):
			os.mkdir(project_fold_leibie)  # 如果没有该文件夹就创建一个
		project_fold_move = os.path.join(os.path.split(project_fold_leibie)[0], os.path.split(project_fold_name)[1])
		if os.path.isdir(project_fold_move):  # 判断原来项目类别文件夹中是否存在该项目文件，如果存在就删除该文件
			os.rmdir(project_fold_leibie)
		shutil.move(project_fold_name, project_fold_leibie)
		flag_filename2excel = 'ok'
	if flag_filename2excel == 'ok':
	# 通过get_sheet()获取的sheet有write()方法
		book = xlrd.open_workbook(r'E:\教务处工作\19年校级质量工程\result.xls')
		write_book = copy(book)
		write_sheet = write_book.get_sheet(u'文件读取记录表')
		sheet1 = book.sheet_by_name(u'文件读取记录表')
		nrows = sheet1.nrows  # 行数
		write_sheet.write(nrows, 0, file)
		write_sheet.write(nrows, 1, date)
		write_book.save(r'E:\教务处工作\19年校级质量工程\result.xls')

	return '结题或者验收OK'



	'''txt=doc.paragraphs;
	for text_flag in ((0, (len(txt)-1))):
	    if  '单位'in txt[1].text:
		    worksheet.write(k,0 , txt[1].text.replace('单位：', ''))

	for i in range(1, len(table.rows)):  # 从表格第二行开始循环读取表格数据
		for j in range(0, len(table.columns)):
			result = table.cell(i, j).text;
        #worksheet.write('A%s'%str(i+1),result)  #循环写处理后的数据生成的列表
			worksheet.write(k, j+1, result);'''
#workbook.close()
#read_word(Mpath,file,date)

#info = open(info,'r') #读取文件内容
	#print(info.readline()) #使用readline函数得到一条一条的信息，如果使用read获取全部信息亦可；
	#info.close()

