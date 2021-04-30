#http://www.jb51.net/article/70318.htm
#http://www.cnblogs.com/graphics/articles/2953665.html
#http://shouce.jb51.net/python/
import win32com,re
import os
import docx;
import xlrd
import xlwt
import xlsxwriter
from win32com.client import Dispatch, constants
from docx import Document;

word=win32com.client.Dispatch('word.application')
'''
设置Word的可见性visible，默认情况下，你看不到Word程序；然后设置Word的警告信息是否出现，默认也是不出现，这样你在使用python控制Word的时候不会弹出Word的警告信息。
'''
word.displayalerts=0
word.visible=0
countdoc=word.Documents.Count
print(countdoc)
path=r'C:\Users\Administrator\Downloads\2018-2019大型考试安排表0227.docx'
os.listdir(r'C:\Users\Administrator\Desktop\各学院报送课程思政报名表')


workbook = xlsxwriter.Workbook('result.xlsx')  #生成表格
worksheet = workbook.add_worksheet(u'sheet1')   #在文件中创建一个名为TEST的sheet,不加名字默认为sheet1
worksheet.set_column('A:A',20)  #设置第一列宽度为20像素
bold=workbook.add_format({'bold':True}) #设置一个加粗的格式对象

k=0;
for info in os.listdir(r'C:\Users\Administrator\Desktop\各学院报送课程思政报名表'):
	domain = os.path.abspath(r'C:\Users\Administrator\Desktop\各学院报送课程思政报名表') #获取文件夹的路径，此处其实没必要这么写，目的是为了熟悉os的文件夹操作
	info = os.path.join(domain,info) #将路径与文件名结合起来就是每个文件的完整路径
	doc = Document(info)  # 读入文件
	table = doc.tables[0] #获取文件中的表格集
	txt=doc.paragraphs;
	for text_flag in ((0, (len(txt)-1))):
	    if  '单位'in txt[1].text:
		    worksheet.write(k,0 , txt[1].text.replace('单位：', ''))

	for i in range(1, len(table.rows)):  # 从表格第二行开始循环读取表格数据
		for j in range(0, len(table.columns)):
			result = table.cell(i, j).text;
        #worksheet.write('A%s'%str(i+1),result)  #循环写处理后的数据生成的列表
			worksheet.write(k, j+1, result);
	k=k+1;


workbook.close()


#info = open(info,'r') #读取文件内容
	#print(info.readline()) #使用readline函数得到一条一条的信息，如果使用read获取全部信息亦可；
	#info.close()

#测试读取word方法一
doc = Document(path) #读入文件
print('方法一',doc);
tables = doc.tables #获取文件中的表格集
print('方法一',tables);
table = tables[0  ]#获取文件中的第一个表格
for i in range(1,len(table.rows)):#从表格第二行开始循环读取表格数据
    result = table.cell(i,0).text + "" +table.cell(i,1).text+table.cell(i,2).text + table.cell(i,3).text
    #cell(i,0)表示第(i+1)行第1列数据，以此类推
    #print(result);


doc1=word.Documents.Open(path)
print('方法二',doc1);
t=doc1.Tables[0]
print('方法二',t);

'''#print(type(t))
tt=str(t)
#print(type(tt))
#print("")
#分割字符串
strs=tt.split('')
print(strs[5])

i=0
while i<500:
    t=doc.Tables[i]
    tt=str(t)
    #strs=tt.split('')
    #print(strs[5])
    i=i+1'''
doc.Close()
word.Quit()
